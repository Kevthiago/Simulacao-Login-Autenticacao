# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

# Estatística Aplicada à Simulação de Sistemas Digitais
# Este estudo de caso aborda a simulação de um sistema de login com autenticação.
# Os objetivos principais são estimar o tempo médio de autenticação e a taxa de sucesso do processo de login.

# Importar bibliotecas necessárias
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches
sns.set(style="whitegrid")

# Importar função de geração de texto da IA Groq
from groq_Connect import gerar_texto_groq

# ==========================
# PARÂMETROS DA SIMULAÇÃO
# ==========================
np.random.seed(42)
num_simulacoes = 1000
tempo_autenticacao_medio = 2.0
taxa_sucesso = 0.9
lambda_users = 600
sla_tempo_maximo = 5.0

# ==========================
# GERAR SIMULAÇÃO
# ==========================
eventos = []

for simulacao in range(num_simulacoes):
    num_usuarios = np.random.poisson(lambda_users / 60)
    if num_usuarios == 0:
        continue

    tempos_chegada = np.random.exponential(1 / (lambda_users / 60), num_usuarios)
    tempos_autenticacao = np.random.exponential(tempo_autenticacao_medio, num_usuarios)
    resultados_login = np.random.binomial(1, taxa_sucesso, num_usuarios)

    for i in range(num_usuarios):
        eventos.append({
            "tempo_chegada": tempos_chegada[i],
            "tempo_autenticacao": tempos_autenticacao[i],
            "resultado_login": resultados_login[i]
        })

df_simulacao = pd.DataFrame(eventos)
df_simulacao["tempo_total"] = df_simulacao["tempo_chegada"] + df_simulacao["tempo_autenticacao"]

# ==========================
# CÁLCULO DAS MÉTRICAS
# ==========================
total_requisições = len(df_simulacao)
df_simulacao["tempo_total"] = df_simulacao["tempo_chegada"] + df_simulacao["tempo_autenticacao"]
total_tempo = df_simulacao["tempo_total"].max() - df_simulacao["tempo_chegada"].min()

tempos_sucesso = df_simulacao[df_simulacao['resultado_login'] == 1]['tempo_autenticacao']
tempo_medio_autenticacao = tempos_sucesso.mean()

num_sucessos = df_simulacao['resultado_login'].sum()
num_falhas = len(df_simulacao) - num_sucessos
taxa_sucesso_estimada = num_sucessos / len(df_simulacao)
throughput = total_requisições / total_tempo if total_tempo > 0 else 0
num_sla_violations = (df_simulacao['tempo_autenticacao'] > sla_tempo_maximo).sum()

resumo_simulacao = {
    "Total de Requisições": total_requisições,
    "Tempo Total da Simulação (s)": total_tempo,
    "Throughput (requisições/s)": throughput,
    "Número de Sucessos": num_sucessos,
    "Número de Falhas": num_falhas,
    "Taxa de Sucesso Estimada": taxa_sucesso_estimada,
    "Tempo Médio de Autenticação (s)": tempo_medio_autenticacao,
    "Número de Violações de SLA": num_sla_violations
}

# ==========================
# RELATÓRIO PRINCIPAL (IA)
# ==========================
prompt_relatorio = f"""
Você é um assistente especializado em gerar relatórios técnicos profissionais no formato **ABNT (Associação Brasileira de Normas Técnicas, versão 2025)**.

A seguir estão os dados resumidos da simulação:

{resumo_simulacao}

Gere um **relatório completo, formal e estruturado**, seguindo rigorosamente as normas ABNT de apresentação de trabalhos técnicos e científicos.  
O texto deve estar pronto para ser inserido em um documento Word (.DOCX), com formatação adequada (títulos hierarquizados, espaçamento, recuos e parágrafos justificados).

### ESTRUTURA OBRIGATÓRIA DO RELATÓRIO:

1. **Título e Resumo Executivo**  
   - Apresente o título principal do trabalho em letras maiúsculas.  
   - Escreva um resumo técnico e objetivo, descrevendo o contexto da simulação, seus objetivos e as principais métricas observadas.  
   - Use linguagem impessoal e técnica.

2. **Métricas Gerais Interpretadas**  
   - Explique de forma detalhada cada métrica: total de requisições, throughput, tempo total da simulação, entre outras.  
   - Interprete os resultados de maneira analítica, identificando padrões, tendências e possíveis gargalos.  
   - Utilize parágrafos coerentes e bem estruturados, conforme a norma ABNT NBR 14724/2025.

3. **Métricas de Autenticação**  
   - Analise o tempo médio de autenticação, taxa de sucesso, número de falhas e violações de SLA.  
   - Discuta a relevância desses indicadores para o desempenho geral do sistema.  
   - Evite redundâncias e priorize clareza técnica.

4. **Observações e Recomendações Técnicas**  
   - Apresente conclusões objetivas baseadas nas métricas observadas.  
   - Sugira melhorias e boas práticas para otimização do sistema de login.  
   - Redija recomendações técnicas com foco em confiabilidade, escalabilidade e tempo de resposta.  

### REQUISITOS DE FORMATAÇÃO ABNT:
- Títulos principais em negrito e letras maiúsculas.  
- Subtítulos em negrito, apenas com a inicial maiúscula.  
- Parágrafos justificados, espaçamento entre linhas de 1,5.  
- Linguagem técnica, objetiva e impessoal.  
- Margens de 2,5 cm (superior e esquerda) e 2 cm (inferior e direita).  
- Numeração progressiva de seções conforme padrão ABNT (1, 1.1, 1.2, etc.).  
- Use apenas texto corrido; **não use listas com marcadores** no corpo do relatório.  
- Utilize sempre o português formal e técnico.

Não inclua gráficos nem códigos de simulação; eles serão adicionados posteriormente.

Autor do relatório: Kevin Thiago dos Santos  
Estudante de Ciência da Computação
"""

texto_relatorio = gerar_texto_groq(prompt_relatorio)

# ==========================
# CRIAÇÃO DO DOCUMENTO DOCX
# ==========================
documento = Document()
documento.add_heading('Relatório de Simulação do Sistema de Login', level=1)

for linha in texto_relatorio.split('\n'):
    if linha.strip():
        documento.add_paragraph(linha.strip())

# ==========================
# GERAÇÃO E INSERÇÃO DOS GRÁFICOS
# ==========================

# Gráfico 1 – Histograma do Tempo de Autenticação
plt.figure(figsize=(10,6))
sns.histplot(df_simulacao['tempo_autenticacao'], bins=30, kde=True, color='skyblue')
plt.axvline(tempo_autenticacao_medio, color='red', linestyle='--', label='Média Esperada')
plt.axvline(sla_tempo_maximo, color='orange', linestyle='--', label='SLA Máximo')
plt.title('Distribuição do Tempo de Autenticação')
plt.xlabel('Tempo de Autenticação (s)')
plt.ylabel('Número de Requisições')
plt.legend()
plt.tight_layout()
plt.savefig('grafico_tempo_autenticacao.png', dpi=300)
plt.close()

# Gráfico 2 – Boxplot do Tempo de Autenticação
plt.figure(figsize=(8,5))
sns.boxplot(x=df_simulacao['tempo_autenticacao'], color='lightgreen')
plt.axvline(sla_tempo_maximo, color='orange', linestyle='--', label='SLA Máximo')
plt.title('Boxplot do Tempo de Autenticação')
plt.xlabel('Tempo de Autenticação (s)')
plt.legend()
plt.tight_layout()
plt.savefig('grafico_boxplot_tempo_autenticacao.png', dpi=300)
plt.close()

# Gráfico 3 – Sucessos vs Falhas
counts = df_simulacao['resultado_login'].value_counts().rename({1:'Sucesso',0:'Falha'})
plt.figure(figsize=(6,4))
sns.barplot(x=counts.index, y=counts.values, palette=['green','red'])
plt.title('Número de Logins: Sucesso vs Falha')
plt.ylabel('Número de Requisições')
plt.xlabel('Resultado do Login')
plt.tight_layout()
plt.savefig('grafico_sucessos_falhas.png', dpi=300)
plt.close()

# Gráfico 4 – Throughput ao longo do tempo
df_simulacao['tempo_seg'] = df_simulacao['tempo_total'].cumsum()
janela = 50
throughput_series = df_simulacao['tempo_total'].rolling(window=janela).apply(lambda x: janela / (x.max()-x.min()), raw=True)

plt.figure(figsize=(10,6))
plt.plot(throughput_series, color='purple')
plt.title('Throughput (req/s) ao Longo da Simulação')
plt.xlabel('Requisição')
plt.ylabel('Throughput (req/s)')
plt.tight_layout()
plt.savefig('grafico_throughput.png', dpi=300)
plt.close()

# Gráfico 5 – Violações de SLA
sla_flag = df_simulacao['tempo_autenticacao'] > sla_tempo_maximo
plt.figure(figsize=(6,4))
sns.countplot(x=sla_flag.map({True:'Violação SLA', False:'Dentro do SLA'}), palette=['red','green'])
plt.title('Violação de SLA nas Requisições')
plt.ylabel('Número de Requisições')
plt.xlabel('Status da Requisição')
plt.tight_layout()
plt.savefig('grafico_sla_violacoes.png', dpi=300)
plt.close()

# Inserindo os gráficos no DOCX
documento.add_page_break()
documento.add_heading('Análise Visual dos Resultados', level=2)

graficos = [
    ('Distribuição do Tempo de Autenticação', 'grafico_tempo_autenticacao.png'),
    ('Boxplot do Tempo de Autenticação', 'grafico_boxplot_tempo_autenticacao.png'),
    ('Sucesso vs Falha de Login', 'grafico_sucessos_falhas.png'),
    ('Throughput ao Longo da Simulação', 'grafico_throughput.png'),
    ('Violações de SLA', 'grafico_sla_violacoes.png')
]

for titulo, caminho in graficos:
    documento.add_heading(titulo, level=3)
    documento.add_picture(caminho, width=Inches(5.5))
    documento.add_paragraph(' ')

# ==========================
# ANÁLISE DA IA SOBRE OS GRÁFICOS
# ==========================
prompt_analise_graficos = """
Você é um assistente especializado em análise técnica de dados e visualizações, redigindo textos conforme as normas da **ABNT (Associação Brasileira de Normas Técnicas, versão 2025)**.

Analise os gráficos técnicos gerados a partir de uma simulação de login com autenticação de usuários.  
Apresente uma **interpretação clara, formal e objetiva**, adequada para inclusão em um relatório técnico acadêmico.

### Diretrizes de Redação:
- Evite repetir informações já descritas no relatório principal.  
- Foque na **interpretação visual dos resultados**, destacando:
  - Padrões observados nos gráficos.  
  - Comportamentos atípicos (anomalias).  
  - Consistência e estabilidade do sistema.  
  - Insights relevantes sobre desempenho, gargalos e tempo de resposta.  
- Utilize **linguagem técnica e impessoal**, conforme a norma ABNT NBR 14724/2025.  
- Os parágrafos devem ser **justificados**, com espaçamento de 1,5 e sem listas com marcadores.  
- Títulos e subtítulos devem seguir hierarquia formal (por exemplo: “4.1 Análise dos Gráficos de Desempenho”).  
- O texto deve ter coesão e coerência, apresentando conclusões analíticas sobre os resultados visuais.

Não inclua descrições detalhadas dos dados numéricos — concentre-se apenas na **interpretação dos comportamentos visuais representados nos gráficos**.
"""

analise_graficos = gerar_texto_groq(prompt_analise_graficos)

documento.add_page_break()
documento.add_heading('Análise Visual da IA', level=2)
documento.add_paragraph(analise_graficos)

# Salva o documento final
documento.save('relatorio_simulacao_login_completo.docx')

print("✅ Relatório completo gerado com sucesso: relatorio_simulacao_login_completo.docx")
